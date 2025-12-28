"""User-related API routes."""

import io
import logging
from uuid import UUID

from fastapi import APIRouter, File, HTTPException, UploadFile
from pydantic import BaseModel, Field
from sqlalchemy.orm.attributes import flag_modified

from src.api.dependencies import ClaudeDep, DbDep
from src.api.schemas import (
    EmailSender,
    EmailSenderPreferences,
    EmailSenderPreferencesResponse,
    EmailSenderPreferencesUpdate,
    UserCreate,
    UserResponse,
    UserUpdate,
)
from src.config import DEFAULT_JOB_EMAIL_SENDERS
from src.db.repositories.user import UserRepository
from src.integrations.claude.client import get_model_id

logger = logging.getLogger(__name__)

router = APIRouter()


# ============================================================================
# CV Upload Schemas
# ============================================================================


class CVUploadResponse(BaseModel):
    """Response after CV upload."""

    success: bool
    message: str
    text_length: int
    preview: str  # First 500 characters


class CVResponse(BaseModel):
    """Response for getting CV."""

    has_cv: bool
    text_length: int
    preview: str | None = None
    content: str | None = None


# ============================================================================
# AI Response Generation Schemas
# ============================================================================


class QuestionAnswerRequest(BaseModel):
    """Request to generate an answer for a job application question."""

    question: str = Field(..., description="The question to answer")
    job_title: str | None = Field(None, description="Job title for context")
    company: str | None = Field(None, description="Company name for context")
    job_description: str | None = Field(None, description="Job description for context")
    max_words: int = Field(300, description="Target word count for answer")
    tone: str = Field(
        "professional", description="Tone: professional, enthusiastic, conversational"
    )
    save_answer: bool = Field(False, description="Save this answer for future use")


class QuestionAnswerResponse(BaseModel):
    """Response with generated answer."""

    question: str
    answer: str
    word_count: int
    from_cache: bool = False  # If retrieved from saved answers
    saved: bool = False  # If answer was saved


class SavedAnswer(BaseModel):
    """A saved answer for a common question."""

    question_pattern: str  # Pattern to match similar questions
    answer: str
    used_count: int = 0
    last_used: str | None = None


class SavedAnswersResponse(BaseModel):
    """Response with all saved answers."""

    answers: list[SavedAnswer]
    total: int


# ============================================================================
# CV Parsing Functions
# ============================================================================


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from a PDF file."""
    from pypdf import PdfReader

    pdf = PdfReader(io.BytesIO(file_content))
    text_parts = []

    for page in pdf.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)

    return "\n\n".join(text_parts)


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from a DOCX file."""
    from docx import Document

    doc = Document(io.BytesIO(file_content))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


@router.post("/", response_model=UserResponse)
async def create_user(user: UserCreate, db: DbDep):
    """Create a new user profile."""
    repo = UserRepository(db)

    # Check if email already exists
    existing = await repo.get_by_email(user.email)
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")

    new_user = await repo.create(**user.model_dump())
    return new_user


@router.get("/{user_id}", response_model=UserResponse)
async def get_user(user_id: UUID, db: DbDep):
    """Get user profile by ID."""
    repo = UserRepository(db)
    user = await repo.get(user_id)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    return user


@router.patch("/{user_id}", response_model=UserResponse)
async def update_user(user_id: UUID, user: UserUpdate, db: DbDep):
    """Update user profile."""
    repo = UserRepository(db)
    updated = await repo.update(user_id, **user.model_dump(exclude_unset=True))

    if not updated:
        raise HTTPException(status_code=404, detail="User not found")

    return updated


@router.post("/{user_id}/cv", response_model=CVUploadResponse)
async def upload_cv(
    user_id: UUID,
    db: DbDep,
    file: UploadFile = File(..., description="CV file (PDF, DOCX, or TXT)"),
):
    """
    Upload base CV for a user.

    Accepts PDF, DOCX, or TXT files. Extracts text and stores it in the user profile.
    Maximum file size: 5MB.
    """
    # Validate file type
    allowed_types = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
    }

    content_type = file.content_type
    if content_type not in allowed_types:
        # Also check by extension
        filename = file.filename or ""
        ext = filename.lower().split(".")[-1] if "." in filename else ""
        if ext not in ["pdf", "docx", "txt"]:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {content_type}. Please upload PDF, DOCX, or TXT.",
            )
        file_type = ext
    else:
        file_type = allowed_types[content_type]

    # Read file content
    content = await file.read()

    # Check file size (5MB max)
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(
            status_code=400,
            detail="File too large. Maximum size is 5MB.",
        )

    # Extract text based on file type
    try:
        if file_type == "pdf":
            text = extract_text_from_pdf(content)
        elif file_type == "docx":
            text = extract_text_from_docx(content)
        else:  # txt
            text = content.decode("utf-8", errors="replace")
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to parse file: {str(e)}",
        )

    if not text or len(text.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="Could not extract meaningful text from the file. Please check the file content.",
        )

    # Clean up text
    text = text.strip()

    # Get user and update CV
    repo = UserRepository(db)
    user = await repo.get(user_id)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    # Update base_cv_content
    user.base_cv_content = text
    await db.flush()
    await db.refresh(user)

    return CVUploadResponse(
        success=True,
        message=f"CV uploaded successfully ({file_type.upper()} format)",
        text_length=len(text),
        preview=text[:500] + "..." if len(text) > 500 else text,
    )


@router.get("/{user_id}/cv", response_model=CVResponse)
async def get_cv(
    user_id: UUID,
    db: DbDep,
    include_content: bool = False,
):
    """
    Get user's base CV.

    By default, only returns metadata. Set include_content=true to get full text.
    """
    repo = UserRepository(db)
    user = await repo.get(user_id)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    cv_text = user.base_cv_content

    if not cv_text:
        return CVResponse(has_cv=False, text_length=0)

    return CVResponse(
        has_cv=True,
        text_length=len(cv_text),
        preview=cv_text[:500] + "..." if len(cv_text) > 500 else cv_text,
        content=cv_text if include_content else None,
    )


@router.delete("/{user_id}/cv")
async def delete_cv(user_id: UUID, db: DbDep):
    """Delete user's base CV."""
    repo = UserRepository(db)
    user = await repo.get(user_id)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    user.base_cv_content = None
    await db.flush()

    return {"success": True, "message": "CV deleted successfully"}


@router.get("/{user_id}/preferences")
async def get_preferences(user_id: UUID, db: DbDep):
    """Get user job preferences."""
    repo = UserRepository(db)
    user = await repo.get(user_id)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    return user.preferences or {}


@router.put("/{user_id}/preferences")
async def update_preferences(user_id: UUID, preferences: dict, db: DbDep):
    """Update user job preferences."""
    repo = UserRepository(db)
    user = await repo.get(user_id)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    user.preferences = preferences
    flag_modified(user, "preferences")
    await db.flush()
    await db.refresh(user)

    return user.preferences


# ============================================================================
# Email Sender Preferences
# ============================================================================


def _get_default_senders() -> list[EmailSender]:
    """Get default email senders as EmailSender objects."""
    return [
        EmailSender(
            id=s["id"],
            name=s["name"],
            pattern=s["pattern"],
            enabled=s.get("enabled", True),
            is_custom=False,
        )
        for s in DEFAULT_JOB_EMAIL_SENDERS
    ]


def _merge_sender_preferences(
    defaults: list[EmailSender],
    user_prefs: EmailSenderPreferences | None,
) -> list[EmailSender]:
    """Merge default senders with user preferences to get effective list."""
    if not user_prefs:
        return [s for s in defaults if s.enabled]

    effective = []

    for sender in defaults:
        # Check if user has overridden the default enabled state
        if sender.id in (user_prefs.disabled_sender_ids or []):
            continue  # User disabled this default
        elif sender.id in (user_prefs.enabled_sender_ids or []):
            effective.append(sender)  # User enabled this default
        elif sender.enabled:
            effective.append(sender)  # Default is enabled

    # Add user's custom senders
    for custom in user_prefs.senders or []:
        if custom.enabled:
            effective.append(custom)

    return effective


@router.get("/{user_id}/email-senders", response_model=EmailSenderPreferencesResponse)
async def get_email_sender_preferences(user_id: UUID, db: DbDep):
    """Get user's email sender preferences."""
    repo = UserRepository(db)
    user = await repo.get(user_id)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    defaults = _get_default_senders()

    # Extract email preferences from user.preferences JSON
    user_email_prefs = None
    if user.preferences and "email_senders" in user.preferences:
        user_email_prefs = EmailSenderPreferences(**user.preferences["email_senders"])

    effective = _merge_sender_preferences(defaults, user_email_prefs)

    return EmailSenderPreferencesResponse(
        default_senders=defaults,
        user_preferences=user_email_prefs or EmailSenderPreferences(),
        effective_senders=effective,
    )


@router.put("/{user_id}/email-senders", response_model=EmailSenderPreferencesResponse)
async def update_email_sender_preferences(
    user_id: UUID,
    updates: EmailSenderPreferencesUpdate,
    db: DbDep,
):
    """Update user's email sender preferences."""
    repo = UserRepository(db)
    user = await repo.get(user_id)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    # Get current preferences or create new
    current_prefs = dict(user.preferences) if user.preferences else {}
    email_prefs = current_prefs.get(
        "email_senders",
        {
            "senders": [],
            "enabled_sender_ids": [],
            "disabled_sender_ids": [],
        },
    )

    # Track enabled/disabled sender IDs
    enabled_ids = set(email_prefs.get("enabled_sender_ids", []))
    disabled_ids = set(email_prefs.get("disabled_sender_ids", []))
    custom_senders = list(email_prefs.get("senders", []))

    # Apply updates
    if updates.enabled_sender_ids:
        enabled_ids.update(updates.enabled_sender_ids)
        disabled_ids -= set(updates.enabled_sender_ids)

    if updates.disabled_sender_ids:
        disabled_ids.update(updates.disabled_sender_ids)
        enabled_ids -= set(updates.disabled_sender_ids)

    if updates.custom_senders:
        for sender in updates.custom_senders:
            sender_dict = sender.model_dump()
            sender_dict["is_custom"] = True
            # Only add if not already exists
            if not any(s.get("id") == sender.id for s in custom_senders):
                custom_senders.append(sender_dict)

    if updates.remove_sender_ids:
        custom_senders = [s for s in custom_senders if s.get("id") not in updates.remove_sender_ids]

    # Update preferences
    email_prefs["enabled_sender_ids"] = list(enabled_ids)
    email_prefs["disabled_sender_ids"] = list(disabled_ids)
    email_prefs["senders"] = custom_senders

    current_prefs["email_senders"] = email_prefs

    # Save to database - need to flag JSON field as modified for SQLAlchemy
    user.preferences = current_prefs
    flag_modified(user, "preferences")
    await db.flush()
    await db.refresh(user)

    # Return updated state
    return await get_email_sender_preferences(user_id, db)


# ============================================================================
# AI Response Generation Endpoints
# ============================================================================


ANSWER_GENERATION_PROMPT = """You are helping a job applicant write personalized responses for job application questions.

## Applicant's CV/Resume:
{cv_content}

## Job Context:
- Position: {job_title}
- Company: {company}
- Job Description: {job_description}

## Question to Answer:
{question}

## Instructions:
1. Write a {tone} response that highlights relevant experience from the CV
2. Target approximately {max_words} words
3. Be specific and use concrete examples from the CV when possible
4. Tailor the response to the specific company and role
5. Be authentic and avoid generic phrases
6. If the question asks "why this company", research common reasons people want to work there
7. Do NOT include any preamble or explanation - just write the answer directly

Write the answer now:"""


def _find_matching_saved_answer(
    question: str,
    saved_answers: list[dict],
) -> dict | None:
    """Find a saved answer that matches the question pattern."""
    question_lower = question.lower()

    for saved in saved_answers:
        pattern = saved.get("question_pattern", "").lower()
        # Simple pattern matching - check if key phrases match
        if pattern and pattern in question_lower:
            return saved

    return None


@router.post("/{user_id}/generate-answer", response_model=QuestionAnswerResponse)
async def generate_answer(
    user_id: UUID,
    request: QuestionAnswerRequest,
    db: DbDep,
    claude: ClaudeDep,
):
    """
    Generate an AI-powered answer for a job application question.

    Uses the user's CV and job context to create personalized responses.
    Optionally saves the answer for future use with similar questions.
    """
    repo = UserRepository(db)
    user = await repo.get(user_id)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    # Check for cached/saved answer first
    prefs = user.preferences or {}
    saved_answers = prefs.get("saved_answers", [])

    matching = _find_matching_saved_answer(request.question, saved_answers)
    if matching and not request.save_answer:  # Don't use cache if explicitly saving new
        # Update usage stats
        matching["used_count"] = matching.get("used_count", 0) + 1
        matching["last_used"] = __import__("datetime").datetime.utcnow().isoformat()
        user.preferences = prefs
        flag_modified(user, "preferences")
        await db.flush()

        return QuestionAnswerResponse(
            question=request.question,
            answer=matching["answer"],
            word_count=len(matching["answer"].split()),
            from_cache=True,
            saved=True,
        )

    # Get CV content
    cv_content = user.base_cv_content
    if not cv_content:
        raise HTTPException(
            status_code=400,
            detail="No CV found. Please upload your CV first using POST /users/{user_id}/cv",
        )

    # Build prompt
    prompt = ANSWER_GENERATION_PROMPT.format(
        cv_content=cv_content[:4000],  # Limit CV length
        job_title=request.job_title or "Not specified",
        company=request.company or "Not specified",
        job_description=(request.job_description or "Not provided")[:2000],
        question=request.question,
        tone=request.tone,
        max_words=request.max_words,
    )

    # Generate with Claude
    try:
        response = claude.messages.create(
            model=get_model_id(),
            max_tokens=1024,
            messages=[{"role": "user", "content": prompt}],
        )
        answer = response.content[0].text.strip()
    except Exception as e:
        logger.error(f"Claude API error: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate answer: {str(e)}",
        )

    word_count = len(answer.split())

    # Save answer if requested
    saved = False
    if request.save_answer:
        # Create a pattern from the question (simplified version)
        # Remove specific company/job references to make it reusable
        pattern = request.question.lower()
        for word in [request.company, request.job_title]:
            if word:
                pattern = pattern.replace(word.lower(), "")
        pattern = " ".join(pattern.split()[:10])  # First 10 words as pattern

        new_saved = {
            "question_pattern": pattern,
            "answer": answer,
            "used_count": 1,
            "last_used": __import__("datetime").datetime.utcnow().isoformat(),
            "original_question": request.question,
        }

        saved_answers.append(new_saved)
        prefs["saved_answers"] = saved_answers
        user.preferences = prefs
        flag_modified(user, "preferences")
        await db.flush()
        saved = True

    return QuestionAnswerResponse(
        question=request.question,
        answer=answer,
        word_count=word_count,
        from_cache=False,
        saved=saved,
    )


@router.get("/{user_id}/saved-answers", response_model=SavedAnswersResponse)
async def get_saved_answers(user_id: UUID, db: DbDep):
    """Get all saved answers for a user."""
    repo = UserRepository(db)
    user = await repo.get(user_id)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    prefs = user.preferences or {}
    saved = prefs.get("saved_answers", [])

    answers = [
        SavedAnswer(
            question_pattern=s.get("question_pattern", ""),
            answer=s.get("answer", ""),
            used_count=s.get("used_count", 0),
            last_used=s.get("last_used"),
        )
        for s in saved
    ]

    return SavedAnswersResponse(answers=answers, total=len(answers))


@router.delete("/{user_id}/saved-answers/{pattern}")
async def delete_saved_answer(user_id: UUID, pattern: str, db: DbDep):
    """Delete a saved answer by its pattern."""
    repo = UserRepository(db)
    user = await repo.get(user_id)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    prefs = user.preferences or {}
    saved = prefs.get("saved_answers", [])

    # Filter out the matching pattern
    new_saved = [s for s in saved if s.get("question_pattern") != pattern]

    if len(new_saved) == len(saved):
        raise HTTPException(status_code=404, detail="Saved answer not found")

    prefs["saved_answers"] = new_saved
    user.preferences = prefs
    flag_modified(user, "preferences")
    await db.flush()

    return {"success": True, "message": "Saved answer deleted"}
