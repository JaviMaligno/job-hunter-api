"""Document generation service for CV and Cover Letter export."""

import io
from datetime import datetime
from enum import Enum

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fpdf import FPDF
from pydantic import BaseModel, Field


class DocumentFormat(str, Enum):
    """Supported document formats."""

    DOCX = "docx"
    PDF = "pdf"


class DocumentType(str, Enum):
    """Document types."""

    CV = "cv"
    COVER_LETTER = "cover_letter"


class DocumentMetadata(BaseModel):
    """Metadata for document generation."""

    job_title: str | None = None
    company: str | None = None
    candidate_name: str | None = None
    date: str = Field(default_factory=lambda: datetime.now().strftime("%B %d, %Y"))


class DocumentGenerator:
    """Generate professional documents in DOCX and PDF formats."""

    @staticmethod
    def generate_cv_docx(content: str, metadata: DocumentMetadata) -> bytes:
        """Generate a CV document in DOCX format."""
        doc = Document()

        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add header if candidate name provided
        if metadata.candidate_name:
            header = doc.add_heading(metadata.candidate_name, level=0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle with job target
        if metadata.job_title and metadata.company:
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subtitle.add_run(f"CV adapted for {metadata.job_title} at {metadata.company}")
            run.italic = True
            run.font.size = Pt(10)

        doc.add_paragraph()  # Spacing

        # Process content - split by common CV sections
        lines = content.split("\n")
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect section headers (all caps, starts with #, or known sections)
            is_header = (
                line.isupper()
                or line.startswith("#")
                or any(
                    section in line.upper()
                    for section in [
                        "EXPERIENCE",
                        "EDUCATION",
                        "SKILLS",
                        "SUMMARY",
                        "PROFILE",
                        "PROJECTS",
                        "CERTIFICATIONS",
                        "LANGUAGES",
                        "EXPERIENCIA",
                        "EDUCACIÓN",
                        "HABILIDADES",
                        "PERFIL",
                        "PROYECTOS",
                        "CERTIFICACIONES",
                        "IDIOMAS",
                    ]
                )
            )

            if is_header:
                # Clean header text
                header_text = line.lstrip("#").strip()
                doc.add_heading(header_text, level=1)
                current_section = header_text
            elif line.startswith("- ") or line.startswith("• "):
                # Bullet point
                p = doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("  - ") or line.startswith("  • "):
                # Nested bullet
                p = doc.add_paragraph(line[4:], style="List Bullet 2")
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(6)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def generate_cover_letter_docx(content: str, metadata: DocumentMetadata) -> bytes:
        """Generate a cover letter document in DOCX format."""
        doc = Document()

        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # Add date
        date_para = doc.add_paragraph(metadata.date)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()  # Spacing

        # Add recipient info if company provided
        if metadata.company:
            doc.add_paragraph(f"To: Hiring Manager")
            doc.add_paragraph(metadata.company)
            if metadata.job_title:
                doc.add_paragraph(f"Re: {metadata.job_title}")
            doc.add_paragraph()

        # Add content paragraphs
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # Handle single line breaks within paragraphs
                para_text = para_text.replace("\n", " ")
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = 1.15

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def generate_cv_pdf(content: str, metadata: DocumentMetadata) -> bytes:
        """Generate a CV document in PDF format."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)

        # Set fonts - use built-in fonts for reliability
        pdf.set_font("Helvetica", size=11)

        # Add header if candidate name provided
        if metadata.candidate_name:
            pdf.set_font("Helvetica", "B", 18)
            pdf.cell(0, 15, metadata.candidate_name, align="C", new_x="LMARGIN", new_y="NEXT")

        # Add subtitle
        if metadata.job_title and metadata.company:
            pdf.set_font("Helvetica", "I", 10)
            pdf.cell(
                0,
                8,
                f"CV adapted for {metadata.job_title} at {metadata.company}",
                align="C",
                new_x="LMARGIN",
                new_y="NEXT",
            )

        pdf.ln(5)

        # Process content
        pdf.set_font("Helvetica", size=10)
        lines = content.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                pdf.ln(3)
                continue

            # Detect section headers
            is_header = line.isupper() or line.startswith("#") or any(
                section in line.upper()
                for section in [
                    "EXPERIENCE",
                    "EDUCATION",
                    "SKILLS",
                    "SUMMARY",
                    "PROFILE",
                    "PROJECTS",
                    "EXPERIENCIA",
                    "EDUCACIÓN",
                    "HABILIDADES",
                    "PERFIL",
                    "PROYECTOS",
                ]
            )

            if is_header:
                header_text = line.lstrip("#").strip()
                pdf.ln(5)
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_fill_color(240, 240, 240)
                pdf.cell(0, 8, header_text, fill=True, new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("Helvetica", size=10)
            elif line.startswith("- ") or line.startswith("• "):
                pdf.set_x(20)
                pdf.multi_cell(0, 6, f"  {line}")
            else:
                pdf.multi_cell(0, 6, line)

        return bytes(pdf.output())

    @staticmethod
    def generate_cover_letter_pdf(content: str, metadata: DocumentMetadata) -> bytes:
        """Generate a cover letter document in PDF format."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=25)
        pdf.set_margins(25, 25, 25)

        # Date - right aligned
        pdf.set_font("Helvetica", size=11)
        pdf.cell(0, 10, metadata.date, align="R", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(10)

        # Recipient
        if metadata.company:
            pdf.set_font("Helvetica", size=11)
            pdf.cell(0, 6, "To: Hiring Manager", new_x="LMARGIN", new_y="NEXT")
            pdf.cell(0, 6, metadata.company, new_x="LMARGIN", new_y="NEXT")
            if metadata.job_title:
                pdf.set_font("Helvetica", "B", 11)
                pdf.cell(0, 6, f"Re: {metadata.job_title}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(10)

        # Content
        pdf.set_font("Helvetica", size=11)
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip().replace("\n", " ")
            if para_text:
                pdf.multi_cell(0, 7, para_text)
                pdf.ln(5)

        return bytes(pdf.output())

    def generate(
        self,
        content: str,
        format: DocumentFormat,
        doc_type: DocumentType,
        metadata: DocumentMetadata | None = None,
    ) -> bytes:
        """
        Generate a document in the specified format.

        Args:
            content: The text content to include
            format: Output format (docx or pdf)
            doc_type: Type of document (cv or cover_letter)
            metadata: Optional metadata for headers/footers

        Returns:
            Document as bytes
        """
        metadata = metadata or DocumentMetadata()

        if doc_type == DocumentType.CV:
            if format == DocumentFormat.DOCX:
                return self.generate_cv_docx(content, metadata)
            else:
                return self.generate_cv_pdf(content, metadata)
        else:
            if format == DocumentFormat.DOCX:
                return self.generate_cover_letter_docx(content, metadata)
            else:
                return self.generate_cover_letter_pdf(content, metadata)
